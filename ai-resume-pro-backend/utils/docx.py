import json
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_docx(content: str) -> str:
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")

    try:
        clean = content.replace("```json", "").replace("```", "").strip()
        data = json.loads(clean)
        return _create_structured_docx(temp_file.name, data)
    except Exception:
        return _create_plain_docx(temp_file.name, content)


def _create_structured_docx(path: str, d: dict) -> str:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Name
    name_para = doc.add_paragraph()
    run = name_para.add_run(d.get("name", ""))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Role
    role_para = doc.add_paragraph()
    run = role_para.add_run(d.get("role", ""))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)

    # Contacts
    contacts = " · ".join(filter(None, [d.get("email"), d.get("phone"), d.get("linkedin")]))
    if contacts:
        cp = doc.add_paragraph()
        run = cp.add_run(contacts)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("─" * 70)

    def add_section_heading(title):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    if d.get("summary"):
        add_section_heading("PROFILE")
        doc.add_paragraph(d["summary"])

    if d.get("experience"):
        add_section_heading("EXPERIENCE")
        for exp in d["experience"]:
            p = doc.add_paragraph()
            r = p.add_run(exp.get("title", ""))
            r.bold = True
            r.font.size = Pt(11)

            cp = doc.add_paragraph()
            r2 = cp.add_run(f"{exp.get('company', '')} · {exp.get('duration', '')}")
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)
            r2.italic = True

            for pt in exp.get("points", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(pt).font.size = Pt(10)

    if d.get("skills"):
        add_section_heading("SKILLS")
        skills = " · ".join(d["skills"]) if isinstance(d["skills"], list) else d["skills"]
        doc.add_paragraph(skills)

    if d.get("education"):
        add_section_heading("EDUCATION")
        doc.add_paragraph(d["education"])

    doc.save(path)
    return path


def _create_plain_docx(path: str, content: str) -> str:
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)
    return path
